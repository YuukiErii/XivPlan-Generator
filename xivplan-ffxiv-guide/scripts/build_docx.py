#!/usr/bin/env python3
"""Build a DOCX guide from guide.json."""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def read_json(path: Path) -> Any:
    return json.loads(path.read_text(encoding="utf-8"))


def set_run_font(run: Any, font_name: str = "Microsoft YaHei", size: int | None = None) -> None:
    run.font.name = font_name
    if size is not None:
        run.font.size = Pt(size)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def add_paragraph(doc: Document, text: str = "", style: str | None = None) -> Any:
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_heading("", level=level)
    run = paragraph.add_run(text)
    set_run_font(run, size=22 if level == 1 else 16)


def add_bullets(doc: Document, items: list[Any]) -> None:
    if not items:
        add_paragraph(doc, "无", "List Bullet")
        return
    for item in items:
        add_paragraph(doc, str(item), "List Bullet")


def add_role_table(doc: Document, assignments: list[Any]) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = table.rows[0].cells
    headers[0].text = "职能"
    headers[1].text = "位置"
    headers[2].text = "任务"
    for item in assignments or [{"role": "全员", "position": "按图示", "task": "按攻略正文处理"}]:
        cells = table.add_row().cells
        cells[0].text = str(item.get("role", ""))
        cells[1].text = str(item.get("position", ""))
        cells[2].text = str(item.get("task", ""))


def add_image_if_present(doc: Document, image_path: Path, caption: str) -> None:
    if not image_path.exists():
        add_paragraph(doc, f"图片未找到：{image_path}")
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(5.8))
    caption_paragraph = add_paragraph(doc, caption)
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_short_callout(doc: Document, guide: dict[str, Any]) -> None:
    add_heading(doc, f"{guide['title']} 队内速记", 1)
    add_bullets(doc, guide.get("short_callout", []))
    add_heading(doc, "职能分工", 2)
    add_role_table(doc, guide.get("role_assignments", []))
    add_heading(doc, "口诀", 2)
    add_paragraph(doc, str(guide.get("mnemonic", "无")))


def build_docx_from_guide(guide: dict[str, Any], output_path: Path, base_dir: Path | None = None, short_only: bool = False) -> None:
    base_dir = base_dir or output_path.parent
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    if short_only:
        add_short_callout(doc, guide)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        return

    add_heading(doc, guide["title"], 1)
    add_heading(doc, "一句话概括", 2)
    add_paragraph(doc, guide["summary"])
    if guide.get("recommended_solution"):
        add_heading(doc, "推荐解法", 2)
        add_paragraph(doc, str(guide["recommended_solution"]))

    for index, figure in enumerate(guide.get("figures", []), start=1):
        title = figure.get("title", f"图 {index}")
        add_heading(doc, f"图 {index}：{title}", 2)
        image = figure.get("image")
        if image:
            add_image_if_present(doc, base_dir / image, figure.get("caption", title))
        add_paragraph(doc, str(figure.get("guide_text") or figure.get("caption") or ""))

    flow = guide.get("flow", [])
    if flow:
        add_heading(doc, "详细处理流程", 2)
        for item in flow:
            add_paragraph(doc, str(item), "List Number")

    add_heading(doc, "职能分工", 2)
    add_role_table(doc, guide.get("role_assignments", []))
    add_heading(doc, "常见失误", 2)
    add_bullets(doc, guide.get("common_mistakes", []))
    add_heading(doc, "队内速记版", 2)
    add_bullets(doc, guide.get("short_callout", []))
    add_heading(doc, "口诀", 2)
    add_paragraph(doc, str(guide.get("mnemonic", "无")))
    add_heading(doc, "图文一致性检查", 2)
    add_bullets(doc, guide.get("consistency_checks", []))
    add_heading(doc, "需要确认的点", 2)
    add_bullets(doc, guide.get("unknowns", []) or ["无"])

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> int:
    parser = argparse.ArgumentParser(description="Build guide.docx from guide.json.")
    parser.add_argument("guide", type=Path, help="Path to guide.json")
    parser.add_argument("-o", "--output", type=Path, required=True, help="Output DOCX path")
    parser.add_argument("--short-only", action="store_true", help="Render only the team-callout version")
    args = parser.parse_args()

    try:
        build_docx_from_guide(read_json(args.guide), args.output, base_dir=args.guide.parent, short_only=args.short_only)
    except Exception as exc:  # noqa: BLE001 - CLI should surface document-generation errors.
        print(f"ERROR: {exc}", file=sys.stderr)
        return 1
    print(f"Wrote {args.output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
